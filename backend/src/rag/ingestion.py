from __future__ import annotations
import hashlib
from pathlib import Path
from dataclasses import dataclass, field
from src.utils.logging import get_logger
from src.utils.text import clean_text

logger = get_logger(__name__)


@dataclass
class RawDocument:
    content: str
    metadata: dict = field(default_factory=dict)
    doc_id: str = ""

    def __post_init__(self) -> None:
        if not self.doc_id:
            self.doc_id = hashlib.sha256(self.content[:512].encode()).hexdigest()[:16]


async def load_pdf(path: Path) -> RawDocument:
    from src.utils.pdf import extract_text_from_pdf
    content = extract_text_from_pdf(path)
    return RawDocument(
        content=clean_text(content),
        metadata={
            "source": str(path),
            "title": path.stem,
            "file_type": "pdf",
        },
    )


async def load_text(path: Path) -> RawDocument:
    content = path.read_text(encoding="utf-8", errors="replace")
    return RawDocument(
        content=clean_text(content),
        metadata={
            "source": str(path),
            "title": path.stem,
            "file_type": path.suffix.lstrip(".") or "txt",
        },
    )


async def load_url(url: str) -> RawDocument:
    import httpx
    from bs4 import BeautifulSoup

    headers = {"User-Agent": "Lumora/1.0 (research assistant)"}
    async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
        response = await client.get(url, headers=headers)
        response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")
    # Remove noise
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    title = soup.find("title")
    title_text = title.get_text(strip=True) if title else url

    main = soup.find("main") or soup.find("article") or soup.body or soup
    content = clean_text(main.get_text(separator="\n"))

    return RawDocument(
        content=content,
        metadata={
            "source": url,
            "title": title_text,
            "file_type": "url",
        },
    )


async def load_docx(path: Path) -> RawDocument:
    import docx
    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    content = "\n\n".join(paragraphs)
    return RawDocument(
        content=clean_text(content),
        metadata={
            "source": str(path),
            "title": path.stem,
            "file_type": "docx",
        },
    )


async def load_document(source: str, file_type: str | None = None) -> RawDocument:
    path = Path(source)
    if source.startswith("http://") or source.startswith("https://"):
        return await load_url(source)
    elif path.suffix.lower() == ".pdf":
        return await load_pdf(path)
    elif path.suffix.lower() in (".docx", ".doc"):
        return await load_docx(path)
    else:
        return await load_text(path)
